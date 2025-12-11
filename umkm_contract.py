import streamlit as st
from datetime import datetime
from io import BytesIO
from docx import Document

st.set_page_config(page_title="Generator Kontrak Kerja UMKM — Tabs", layout="wide")


def fmt_date(d):
    try:
        return d.strftime("%d %B %Y")
    except Exception:
        return str(d)


# ======================
# TEMPLATES
# ======================
TEMPLATE_PKWTT = """PERJANJIAN KERJA WAKTU TIDAK TERTENTU (PKWTT)

Pada hari ini, {today}, yang bertanda tangan di bawah ini:

Pihak I (Pemberi Kerja):
- Nama usaha: {company_name}
- Alamat: {company_address}
- Nama penanggung jawab: {employer_name}, Jabatan: {employer_position}

Pihak II (Karyawan):
- Nama: {employee_name}
- NIK / KTP: {employee_id}
- Alamat: {employee_address}

Pasal 1 — Jabatan dan Tanggung Jawab
Pihak II diangkat sebagai {position} dan bertanggung jawab melakukan: {job_scope}. Pihak II wajib melaksanakan tugas berdasarkan instruksi atasan dan standar operasional perusahaan.

Pasal 2 — Mulai Bekerja & Status
1. Pihak II mulai bekerja pada tanggal: {start_date}.
2. Perjanjian ini bersifat tidak terbatas waktu (permanent) dan berlaku sejak tanggal mulai kerja, kecuali diakhiri sesuai ketentuan dalam perjanjian ini dan peraturan perundang-undangan.

Pasal 3 — Masa Percobaan
1. Masa percobaan ditetapkan selama {probation_months} bulan sejak tanggal dimulai.
2. Selama masa percobaan, Pihak I akan melakukan evaluasi berdasarkan kriteria: kompetensi teknis, kedisiplinan, kemampuan berkomunikasi, dan pencapaian KPI.
3. Apabila tidak lulus masa percobaan, hubungan kerja berakhir sesuai ketentuan, dengan pemberitahuan tertulis.

Pasal 4 — Upah, Tunjangan, dan Fasilitas
1. Upah pokok: Rp {salary} per bulan.
2. Tunjangan: {allowances}.
3. Penggajian dilakukan setiap tanggal {payday} melalui {payment_method}.
4. Fasilitas yang diinginkan Pihak II: {facilities}.

Pasal 5 — Jam Kerja, Lembur, dan Cuti
1. Jam kerja: {work_hours}. 
2. Hari kerja: {work_days}.
3. Cuti tahunan: {annual_leave} hari per tahun.
4. Lembur dibayar sesuai ketentuan perundang-undangan dan kebijakan perusahaan: persetujuan atasan harus diperoleh sebelum lembur.

Pasal 6 — Hak dan Kewajiban
1. Hak pihak II: menerima upah tepat waktu, cuti tahunan, cuti sakit, cuti melahirkan, dan cuti lainnya mengikuti ketentuan yang berlaku, fasilitas sesuai kebijakan.
2. Kewajiban pihak II: melaksanakan tugas dengan baik, lembur jika diperlukan, mematuhi tata tertib, menjaga aset dan kerahasiaan.

Pasal 7 — Disiplin dan Sanksi
1. Perusahaan memiliki peraturan disiplin internal. Pelanggaran akan dikenai sanksi administratif sampai pemutusan hubungan kerja (PHK) bila termasuk pelanggaran berat.
2. PHK dilakukan sesuai ketentuan perundang-undangan dan kebijakan internal. Pihak yang ingin mengakhiri hubungan kerja wajib memberikan pemberitahuan tertulis {notice_days} hari kecuali PHK karena pelanggaran berat.
3. PHK oleh Pihak I dapat dilakukan sesuai alasan objektif (restrukturisasi, pelanggaran, efisiensi) dengan mekanisme sesuai peraturan ketenagakerjaan.
4. Kompensasi, pesangon dan hak-hak lain dihitung sesuai masa kerja dan perundang-undangan yang berlaku.

Pasal 8 — Kepemilikan Hasil Kerja
Semua hasil kerja yang dibuat Pihak II terkait pekerjaan termasuk namun tidak terbatas pada dokumen, resep, desain, kode, menjadi hak milik Pihak I kecuali disepakati lain.

Pasal 9 — Penyelesaian Perselisihan
Setiap perselisihan diselesaikan secara musyawarah; jika tidak tercapai, dapat ditempuh mediasi/putusan pengadilan sesuai ketentuan hukum yang berlaku.

Pasal 10 — Force Majeure
Kedua pihak tidak bertanggung jawab atas keterlambatan atau kegagalan pelaksanaan yang disebabkan oleh kejadian di luar kemampuan (bencana, kebijakan pemerintah, pandemi) selama kejadian tersebut berlangsung.

Pasal 11 — Kerjasama & Perubahan Kontrak
Perubahan terhadap perjanjian ini harus dibuat tertulis dan ditandatangani kedua pihak.

Tanda tangan
Pihak I: ______________________

Pihak II: ______________________

Saksi: ______________________
"""

TEMPLATE_PKWT = """PERJANJIAN KERJA WAKTU TERTENTU (PKWT)

Pada hari ini, {today}, yang bertanda tangan di bawah ini:

Pihak I (Pemberi Kerja):
- Nama usaha: {company_name}
- Alamat: {company_address}
- Nama penanggung jawab: {employer_name}, Jabatan: {employer_position}

Pihak II (Pekerja):
- Nama: {employee_name}
- NIK / KTP: {employee_id}
- Alamat: {employee_address}

Pasal 1 — Ruang Lingkup Pekerjaan
Pihak II dipekerjakan sebagai {position} dengan uraian tugas utama: {job_scope}. Pekerjaan ini bersifat sementara/proyek: {pkwt_reason}.

Pasal 2 — Masa Perjanjian
1. Perjanjian ini mulai berlaku pada tanggal {start_date} sampai dengan tanggal {end_date}, kecuali diakhiri lebih awal sesuai ketentuan pada perjanjian ini.
2. Perpanjangan hanya dapat dilakukan bila ada kesepakatan tertulis sebelum masa berakhir.

Pasal 3 — Upah dan Tunjangan
1. Upah pokok: Rp {salary} per {pay_period}.
2. Metode pembayaran: {payment_method}, setiap {payday}.
3. Tunjangan (jika ada): {allowances}.

Pasal 4 — Cuti & Pengganti
1. Karena status kontrak sementara, hak cuti dihitung proporsional sesuai kebijakan.
2. Jika pekerjaan terhenti sebelum jangka waktu berakhir bukan karena pelanggaran Pihak II, disepakati kompensasi proporsional.

Pasal 5 — Jam Kerja dan Hari Kerja
Jam kerja: {work_hours}, istirahat {break_minutes} menit. Hari kerja: {work_days}.

Pasal 6 — Pemutusan dan Pengakhiran
1. Perjanjian berakhir otomatis pada tanggal berakhirnya.
2. Pemutusan lebih awal dapat dilakukan jika salah satu pihak melanggar ketentuan material setelah mendapat peringatan tertulis.
3. Kompensasi akhir dihitung sesuai masa kerja dan ketentuan yang berlaku.

Pasal 7 — Kerahasiaan
Pihak II wajib menjaga kerahasiaan informasi usaha yang bersifat strategis dan tidak dibocorkan kepada pihak ketiga selama dan setelah berakhirnya perjanjian.

Pasal 8 — Lain-lain
Hal-hal yang belum diatur akan diselesaikan secara musyawarah dan mengacu pada peraturan perundang-undangan yang berlaku.

Tanda tangan
Pihak I: ______________________

Pihak II: ______________________

Saksi: ______________________
"""

TEMPLATE_PARTTIME = """PERJANJIAN KERJA PART-TIME (PARUH WAKTU)

Pada hari ini, {today}, antara:

Pihak I (Pemberi Kerja): {company_name}

Pihak II (Pekerja Part-Time): {employee_name}, NIK: {employee_id}

Pasal 1 — Bentuk Pekerjaan
Pihak II dipekerjakan secara paruh waktu untuk pekerjaan: {job_scope}. Jam kerja disepakati sebagai berikut: {work_schedule}.

Pasal 2 — Hak & Kewajiban
1. Upah: Rp {salary} per {pay_period}. Pembayaran dilakukan setiap {payday} melalui {payment_method}. Hak-hak (cuti, tunjangan) disesuaikan proporsional. Tidak semua tunjangan penuh berlaku untuk part-time; hak proporsional (mis. cuti) akan dihitung sesuai persentase jam kerja terhadap full-time.
2. Pihak II wajib menjalankan tugas sesuai uraian, menjaga etika, dan melaporkan ketidakhadiran sesuai prosedur. 

Pasal 3 — Durasi dan Jadwal
Perjanjian berlaku sejak {start_date} sampai {end_date} atau berdasarkan kesepakatan jadwal kerja mingguan.

Pasal 4 — Overtime & Permintaan Tambahan
1. Kerja di luar jadwal normal harus disetujui terlebih dahulu; kompensasi tambahannya akan dibayar sesuai kebijakan.
2. Pihak I berhak meminta penyesuaian jam sewaktu-waktu berdasarkan kebutuhan operasional dengan kompensasi dan pemberitahuan sesuai kesepakatan.

Pasal 5 — Pemutusan
Pemutusan dapat dilakukan dengan pemberitahuan tertulis {notice_days} hari kecuali karena pelanggaran berat.

Pasal 6 — Kerahasiaan & Perilaku
Pihak II harus menjaga kerahasiaan data pelanggan dan proses internal selama dan setelah bekerja selama waktu yang disepakati.

Tanda tangan
Pihak I: ______________________

Pihak II: ______________________
"""

TEMPLATE_MAGANG = """PERJANJIAN MAGANG / KONTRAK MAGANG

Pada hari ini, {today}, antara:

Pihak I (Penyelenggara/Mentor): {company_name}

Pihak II (Peserta Magang): {employee_name}

Pasal 1 — Tujuan
Magang ini bertujuan memberikan pengalaman kerja, penguasaan kompetensi dan pembelajaran praktis terkait: {job_scope}.

Pasal 2 — Durasi & Jam Kerja
Durasi magang: {probation_months} bulan (mulai {start_date} sampai {end_date}). Jam kerja: {intern_schedule}.

Pasal 3 — Pembimbing & Evaluasi
Pihak I menunjuk pembimbing untuk membimbing, menilai, dan memberikan sertifikat bila lulus.

Pasal 4 — Imbalan / Stipend
Jika disepakati, peserta mendapat stipend sebesar Rp {salary} per {pay_period} atau fasilitas lain. Peserta tidak dianggap sebagai karyawan; hak ketenagakerjaan penuh tidak otomatis berlaku kecuali disesuaikan.

Pasal 5 — Hak & Kewajiban
1. Hak: pembelajaran, bimbingan, akses sumber belajar, dan sertifikat/ surat keterangan magang apabila dinyatakan lulus. 
2. Kewajiban: menjalankan tugas sesuai rencana pembelajaran, menjaga kerahasiaan, menghormati etika organisasi, dan menyelesaikan laporan magang.

Pasal 6 — Kerahasiaan & Publikasi
Data dan informasi internal yang diakses peserta adalah rahasia. Publikasi hasil magang (contoh portofolio) memerlukan persetujuan tertulis. Kepemilikan hasil kerja untuk tujuan komersial harus disepakati secara tertulis.

Pasal 7 — Pengakhiran Program
Program dapat dihentikan lebih awal jika peserta melakukan pelanggaran, tidak memenuhi persyaratan, atau karena force majeure. Pengakhiran dicatat secara tertulis dengan alasan dan dokumentasi.

Tanda tangan
Pihak I: ______________________

Pihak II: ______________________
"""

TEMPLATE_FREELANCE = """PERJANJIAN KERJA FREELANCE / KERJA LEPAS

Pada hari ini, {today}, antara:

Pemberi Tugas (Klien): {company_name}

Penyedia Jasa (Freelancer): {employee_name}

Pasal 1 — Ruang Lingkup & Deliverables
Freelancer menyelesaikan: {job_scope} sesuai deliverables dan jadwal yang disepakati.

Pasal 2 — Durasi & Jadwal
Waktu pelaksanaan: mulai {start_date} sampai {end_date}. Keterlambatan penyampaian yang tidak dibenarkan berpotensi dikenai penalti atau pemotongan pembayaran setelah evaluasi.

Pasal 2 — Imbalan & Pembayaran
Imbalan: Rp {salary} per {pay_period} atau sesuai kesepakatan. Mekanisme pembayaran dan pajak diatur sesuai kesepakatan. Freelancer bertanggung jawab terhadap kewajiban pajak penghasilannya; Pemberi Tugas tidak memotong hak ketenagakerjaan.

Pasal 4 — Revisi & Acceptance
Paket revisi: [JUMLAH] kali revisi diperbolehkan sesuai cakupan; revisi tambahan dikenakan biaya.
Setelah penyampaian, Pemberi Tugas berwenang melakukan acceptance test dalam [X hari]. Jika tidak ada keberatan tertulis, deliverable dianggap diterima.

Pasal 3 — IP & Kepemilikan
Kecuali disepakati lain, setelah pembayaran penuh, hak atas deliverable dialihkan kepada Pemberi Tugas. Freelancer berhak disebut sebagai pembuat/kreator kecuali ada klausul pengalihan nama.

Pasal 4 — Kerahasiaan
Freelancer wajib menjaga kerahasiaan informasi sensitif dan tidak menyebarkan data klien.
Periode kerahasiaan: selama kerja dan {nda_years} setelah kontrak berakhir (atau sesuai kesepakatan).

Pasal 9 — Pembatalan & Pengakhiran
Kedua pihak dapat mengakhiri kontrak dengan pemberitahuan tertulis dan menyelesaikan kewajiban proporsional.
Jika pengakhiran karena kelalaian Freelancer, Pemberi Tugas bisa menuntut penggantian biaya sesuai bukti.

Pasal 10 — Force Majeure & Penyelesaian Sengketa
Keadaan kahar membebaskan pihak dari kewajiban sementara; jadwal dan biaya akan dinegosiasikan ulang bila perlu.
Perselisihan diselesaikan melalui musyawarah, mediasi, atau arbitrase/pengadilan sesuai perjanjian.

Tanda tangan
Pemberi Tugas: ______________________

Penyedia Jasa: ______________________
"""


# ======================
# HELPERS
# ======================
def base_context(
    company_name,
    company_address,
    employer_name,
    employer_position,
    employee_name,
    employee_id,
    employee_address,
    position,
    job_scope,
    salary_display,
    allowances,
    facilities,
    payment_method,
    pay_period,
    payday,
    start_date,
    end_date,
):
    """Context dasar untuk semua template, dengan default aman untuk field yang tidak dipakai."""
    ctx = {
        "today": fmt_date(datetime.today()),
        "company_name": company_name,
        "company_address": company_address,
        "employer_name": employer_name,
        "employer_position": employer_position,
        "employee_name": employee_name,
        "employee_id": employee_id,
        "employee_address": employee_address,
        "position": position,
        "job_scope": job_scope,
        "salary": salary_display,
        "allowances": allowances,
        "facilities": facilities,
        "payment_method": payment_method,
        "pay_period": pay_period,
        "payday": payday,
        "start_date": fmt_date(start_date),
        "end_date": fmt_date(end_date),
        # default untuk field lain
        "work_hours": "",
        "annual_leave": "",
        "break_minutes": "",
        "work_days": "",
        "work_schedule": "",
        "intern_schedule": "",
        "notice_days": 0,
        "pkwt_reason": "",
        "probation_months": 0,
        "nda_years": 0,
    }
    return ctx


def create_docx(text):
    doc = Document()
    for block in text.split("\n\n"):
        if block.strip() == "" or block.strip() == "---":
            continue
        first_line = block.strip().splitlines()[0]
        if first_line.upper().startswith(("PERJANJIAN", "LAMPIRAN", "PERSETUJUAN")):
            doc.add_heading(first_line.strip(), level=2)
            rest = "\n".join(block.strip().splitlines()[1:]).strip()
            if rest:
                doc.add_paragraph(rest)
        else:
            doc.add_paragraph(block)
    doc.add_paragraph(
        "\n\n\nTanda tangan:\n\nPihak I: ______________________\n\nPihak II: ______________________\n\nSaksi: ______________________"
    )
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ======================
# MAIN PAGE
# ======================
st.title("Generator Template Kontrak Kerja — UMKM (Tab-based)")
st.caption("Pilih jenis kontrak lewat tab di bawah, isi data, lalu generate & download.")


# 1) DATA UMUM (BERLAKU UNTUK SEMUA KONTRAK)
st.header("Data Umum")

col1, col2 = st.columns(2)

with col1:
    company_name = st.text_input("Nama Usaha", value="PT Contoh Usaha")
    company_address = st.text_area("Alamat Usaha", value="Jl. Contoh No.1, Jakarta")
    employer_name = st.text_input("Nama Penanggung Jawab (Pemilik/Manager)", value="Budi Santoso")
    employer_position = st.text_input("Jabatan Penanggung Jawab", value="Pemilik")

with col2:
    employee_name = st.text_input("Nama Karyawan/Pekerja", value="Siti Aminah")
    employee_id = st.text_input("NIK / No. KTP", value="3171xxxxxxxxxxxx")
    employee_address = st.text_area("Alamat Karyawan", value="Jl. Karyawan No.2")

position = st.text_input("Jabatan/Posisi", value="Kasir")
job_scope = st.text_area("Uraian Tugas / Lingkup Pekerjaan", value="Melayani kasir, menghitung kas harian, melayani pelanggan")

salary = st.text_input("Upah / Imbalan (angka tanpa pemisah, mis. 3500000)", value="3500000")
salary_display = f"{int(salary):,}" if salary.isdigit() else salary
allowances = st.text_input("Tunjangan (jika ada)", value="Transport Rp 200.000 / bulan")
facilities = st.text_area("Fasilitas yang Anda inginkan (jika ada)", value="BPJS Kesehatan, BPJS Ketenagakerjaan, Makan Siang, akomodasi")
payment_method = st.selectbox("Metode Pembayaran", ["Transfer Bank", "Tunai"], index=0)
pay_period = st.selectbox("Periode Bayar", ["Bulan", "Minggu", "Hari", "Per Proyek/Deliverable"], index=0)
payday = st.text_input("Tanggal/Jadwal Bayar (mis. setiap 25 / setelah proyek selesai)", value="Setiap tanggal 25")

col3, col4 = st.columns(2)
with col3:
    start_date = st.date_input("Tanggal Mulai", value=datetime.today())
with col4:
    end_date = st.date_input("Tanggal Berakhir (PKWT/Part-Time/Magang/Freelance, opsional)", value=datetime.today())

base_ctx = base_context(
    company_name,
    company_address,
    employer_name,
    employer_position,
    employee_name,
    employee_id,
    employee_address,
    position,
    job_scope,
    salary_display,
    allowances,
    facilities,
    payment_method,
    pay_period,
    payday,
    start_date,
    end_date,
)

st.markdown("---")

# 2) TABS UNTUK TIAP JENIS KONTRAK
tab_pkwtt, tab_pkwt, tab_pt, tab_magang, tab_freelance = st.tabs(
    ["PKWTT (Tetap)", "PKWT (Kontrak)", "Part-Time", "Magang", "Freelance"]
)

# --- TAB PKWTT ---
with tab_pkwtt:
    st.subheader("PKWTT (Perjanjian Kerja Waktu Tidak Tertentu)")

    with st.form("form_pkwtt"):
        probation_months = st.number_input("Masa Percobaan (bulan)", min_value=0, value=3, step=1)
        notice_days = st.number_input("Notice period (hari) untuk PHK", min_value=0, value=30, step=1)
        annual_leave = st.number_input("Cuti Tahunan (hari)", min_value=0, value=12, step=1)
        submitted_pkwtt = st.form_submit_button("Generate Kontrak PKWTT")
        

    if submitted_pkwtt:
        ctx = dict(base_ctx)
        ctx["probation_months"] = probation_months
        ctx["notice_days"] = notice_days
        ctx["annual_leave"] = annual_leave

        text = TEMPLATE_PKWTT.format(**ctx)

        st.markdown("### Pratinjau Kontrak PKWTT")
        st.text_area("Hasil Kontrak PKWTT", value=text, height=420)

        st.write("#### Download")
        txt_bytes = text.encode("utf-8")
        st.download_button(
            "Download PKWTT (.txt)",
            data=txt_bytes,
            file_name="kontrak_pkwtt_umkm.txt",
            mime="text/plain",
            key="dl_pkwtt_txt",
        )
        docx_io = create_docx(text)
        st.download_button(
            "Download PKWTT (.docx)",
            data=docx_io,
            file_name="kontrak_pkwtt_umkm.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_pkwtt_docx",
        )

# --- TAB PKWT ---
with tab_pkwt:
    st.subheader("PKWT (Perjanjian Kerja Waktu Tertentu)")

    with st.form("form_pkwt"):
        work_hours = st.text_input("Jam Kerja", value="08.00–17.00")
        break_minutes = st.text_input("Istirahat (menit)", value="60")
        work_days = st.text_input("Hari Kerja", value="Senin–Jumat")
        pkwt_reason = st.text_input("Alasan PKWT (musiman/proyek/kontrak)", value="Musiman / Proyek")
        notice_days_pkwt = st.number_input("Notice period (hari) untuk pemutusan lebih awal", min_value=0, value=30, step=1)
        submitted_pkwt = st.form_submit_button("Generate Kontrak PKWT")

    if submitted_pkwt:
        ctx = dict(base_ctx)
        ctx["work_hours"] = work_hours
        ctx["break_minutes"] = break_minutes
        ctx["work_days"] = work_days
        ctx["pkwt_reason"] = pkwt_reason
        ctx["notice_days"] = notice_days_pkwt  # jika mau digunakan di pengembangan lanjut

        text = TEMPLATE_PKWT.format(**ctx)

        st.markdown("### Pratinjau Kontrak PKWT")
        st.text_area("Hasil Kontrak PKWT", value=text, height=420)

        st.write("#### Download")
        txt_bytes = text.encode("utf-8")
        st.download_button(
            "Download PKWT (.txt)",
            data=txt_bytes,
            file_name="kontrak_pkwt_umkm.txt",
            mime="text/plain",
            key="dl_pkwt_txt",
        )
        docx_io = create_docx(text)
        st.download_button(
            "Download PKWT (.docx)",
            data=docx_io,
            file_name="kontrak_pkwt_umkm.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_pkwt_docx",
        )

# --- TAB PART-TIME ---
with tab_pt:
    st.subheader("Kontrak Part-Time (Paruh Waktu)")

    with st.form("form_parttime"):
        work_schedule = st.text_input(
            "Jadwal Kerja Part-Time",
            value="Setiap Senin–Jumat pukul 08.00–12.00",
        )
        notice_days_pt = st.number_input("Notice period (hari) untuk pemutusan", min_value=0, value=7, step=1)
        submitted_pt = st.form_submit_button("Generate Kontrak Part-Time")

    if submitted_pt:
        ctx = dict(base_ctx)
        ctx["work_schedule"] = work_schedule
        ctx["notice_days"] = notice_days_pt

        text = TEMPLATE_PARTTIME.format(**ctx)

        st.markdown("### Pratinjau Kontrak Part-Time")
        st.text_area("Hasil Kontrak Part-Time", value=text, height=420)

        st.write("#### Download")
        txt_bytes = text.encode("utf-8")
        st.download_button(
            "Download Part-Time (.txt)",
            data=txt_bytes,
            file_name="kontrak_parttime_umkm.txt",
            mime="text/plain",
            key="dl_pt_txt",
        )
        docx_io = create_docx(text)
        st.download_button(
            "Download Part-Time (.docx)",
            data=docx_io,
            file_name="kontrak_parttime_umkm.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_pt_docx",
        )

# --- TAB MAGANG ---
with tab_magang:
    st.subheader("Kontrak Magang")

    with st.form("form_magang"):
        probation_months_mg = st.number_input("Durasi Magang (bulan)", min_value=1, value=3, step=1)
        intern_schedule = st.text_input("Jadwal Magang", value="Senin–Jumat, 09.00–15.00")
        submitted_mg = st.form_submit_button("Generate Kontrak Magang")

    if submitted_mg:
        ctx = dict(base_ctx)
        ctx["probation_months"] = probation_months_mg
        ctx["intern_schedule"] = intern_schedule

        text = TEMPLATE_MAGANG.format(**ctx)

        st.markdown("### Pratinjau Kontrak Magang")
        st.text_area("Hasil Kontrak Magang", value=text, height=420)

        st.write("#### Download")
        txt_bytes = text.encode("utf-8")
        st.download_button(
            "Download Magang (.txt)",
            data=txt_bytes,
            file_name="kontrak_magang_umkm.txt",
            mime="text/plain",
            key="dl_mg_txt",
        )
        docx_io = create_docx(text)
        st.download_button(
            "Download Magang (.docx)",
            data=docx_io,
            file_name="kontrak_magang_umkm.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_mg_docx",
        )

# --- TAB FREELANCE ---
with tab_freelance:
    st.subheader("Kontrak Freelance / Kerja Lepas")

    with st.form("form_freelance"):
        nda_years_fl = st.number_input("Durasi kerahasiaan (tahun) jika disepakati", min_value=0, value=2, step=1)
        submitted_fl = st.form_submit_button("Generate Kontrak Freelance")

    if submitted_fl:
        ctx = dict(base_ctx)
        ctx["nda_years"] = nda_years_fl

        text = TEMPLATE_FREELANCE.format(**ctx)

        st.markdown("### Pratinjau Kontrak Freelance")
        st.text_area("Hasil Kontrak Freelance", value=text, height=420)

        st.write("#### Download")
        txt_bytes = text.encode("utf-8")
        st.download_button(
            "Download Freelance (.txt)",
            data=txt_bytes,
            file_name="kontrak_freelance_umkm.txt",
            mime="text/plain",
            key="dl_fl_txt",
        )
        docx_io = create_docx(text)
        st.download_button(
            "Download Freelance (.docx)",
            data=docx_io,
            file_name="kontrak_freelance_umkm.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="dl_fl_docx",
        )
